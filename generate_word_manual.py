import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name='TH SarabunPSK', size=16, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

doc = Document()

# ตั้งค่าสไตล์พื้นฐาน
style = doc.styles['Normal']
font = style.font
font.name = 'TH SarabunPSK'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
font.size = Pt(16)

# หน้าปก
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('โรงพยาบาลวังน้ำเย็น จังหวัดสระแก้ว\nกลุ่มงานเภสัชกรรมและคุ้มครองผู้บริโภค\n\n')
set_font(run, size=24, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('คู่มือขั้นตอนการจัดยาและแพ็คยา (OPD / IPD)\nDetailed Dispensing & Packing Procedure\nรหัสเอกสาร SP-003 (ฉบับปรับปรุงใหม่)\n')
set_font(run, size=20, bold=True)

doc.add_page_break()

# ส่วนที่ 1: บทนำ
p = doc.add_paragraph()
run = p.add_run('บทนำ / Introduction')
set_font(run, size=18, bold=True)
p = doc.add_paragraph()
run = p.add_run('คู่มือฉบับนี้จัดทำขึ้นเพื่อกำหนดมาตรฐานขั้นตอนการจัดยาและการบรรจุยา (แพ็คยา) ทั้งผู้ป่วยนอก (OPD) และผู้ป่วยใน (IPD) เพื่อให้เจ้าหน้าที่ผู้จัดยาและเภสัชกรปฏิบัติงานได้อย่างถูกต้อง แม่นยำ ป้องกันความคลาดเคลื่อนทางยา (Medication Error) และคำนึงถึงความปลอดภัยสูงสุดของผู้ป่วย')
set_font(run)

doc.add_paragraph()

# ขั้นตอนต่างๆ
steps = [
    ("ขั้นที่ 1: การรับใบคำสั่งจ่ายยาและการคัดกรอง",
     "• ตรวจสอบความสมบูรณ์ของใบสั่งยาจากแพทย์หรือระบบคอมพิวเตอร์\n"
     "• ตรวจสอบและยืนยันข้อมูลผู้ป่วย (ชื่อ-นามสกุล, HN, หมายเลขห้อง/เตียง)\n"
     "• พิมพ์สติ๊กเกอร์ยาและใบนำทาง พร้อมตรวจสอบความถูกต้องของข้อมูลก่อนพิมพ์"),
     
    ("ขั้นที่ 2: การเตรียมอุปกรณ์และการเลือกซองบรรจุยา (จุดควบคุมคุณภาพการแพ็ค)",
     "• จัดเตรียมตะกร้า ซองยา (ซองใส, ซองชา) และถาดนับยาให้สะอาดและพร้อมใช้งาน\n"
     "• กฎการเลือกซองบรรจุยา:\n"
     "  - ยาทั่วไป: ใช้ซองใส\n"
     "  - ยาไวต่อแสง: ต้องใช้ซองสีชาหรือซองทึบแสงเท่านั้น เพื่อป้องกันยาเสื่อมสภาพ\n"
     "  - ยาที่มีความชื้นหรือต้องการการจัดเก็บพิเศษ: ปฏิบัติตามคำแนะนำของยาเฉพาะเจาะจง"),
     
    ("ขั้นที่ 3: การหยิบยาจากชั้นเก็บยา (ระบบ FEFO)",
     "• อ่านฉลากยาบนสติ๊กเกอร์อย่างละเอียด (ชื่อสามัญ, ชื่อการค้า, ความแรง, จำนวน)\n"
     "• หยิบยาตามหลัก FEFO (First Expired, First Out) เลือกยาที่หมดอายุก่อนออกมาก่อนเสมอ\n"
     "• ตรวจสอบความถูกต้องของยา (ชื่อ, ความแรง) และวันหมดอายุก่อนนำมาจัด"),
     
    ("ขั้นที่ 4: การบรรจุยาลงซอง (ขั้นตอนการแพ็คยา)",
     "• เทยาลงถาดนับยาที่สะอาด ห้ามใช้มือสัมผัสยาโดยตรงเพื่อป้องกันการปนเปื้อน\n"
     "• นับยาให้ครบตามจำนวนที่ระบุ\n"
     "• บรรจุยาลงซองที่เลือกไว้ในขั้นตอนที่ 2\n"
     "  - กรณีเป็นยาแผงฟอยล์: ให้บรรจุทั้งแผง ห้ามตัดหรือแกะเม็ดยาออกจากฟอยล์เอง\n"
     "• ปิดผนึกซองยาให้เรียบร้อยและสนิท เพื่อป้องกันยาหล่นหายหรือความชื้นเข้า"),
     
    ("ขั้นที่ 5: การตรวจสอบและทวนสอบ (Verification & Double Check)",
     "• ผู้จัดยา (Self-Check): ตรวจสอบซองยาทั้งหมดเทียบกับใบสั่งยาอีกครั้งก่อนส่งต่อ\n"
     "• เภสัชกรผู้ทวนสอบ (Final Check): ตรวจสอบยืนยันความถูกต้อง (ประเภทยา, ความแรง, จำนวน, และประเภทของซองยาว่าเหมาะสมหรือไม่)"),
     
    ("ขั้นที่ 6: การจัดคิวและการส่งมอบยาให้ผู้ป่วย",
     "• การจัดเรียงตะกร้ายาตามระบบช่องคิว (ระบบเดียวกับยาฉีด):\n"
     "  - คิว N: คนไข้ทั่วไป\n"
     "  - คิว C: คนไข้โรคเรื้อรัง (ต่อเนื่อง)\n"
     "  - คิว T: คนไข้กลุ่มเร่งด่วน / พระ / เบิกยาด่วน\n"
     "• เภสัชกรเรียกชื่อผู้ป่วยและขอดูบัตรประชาชน/บัตรผู้ป่วยเพื่อยืนยันตัวตน (Identification)\n"
     "• ส่งมอบยาพร้อมให้คำปรึกษา แนะนำวิธีใช้ ผลข้างเคียง และการเก็บรักษา\n"
     "• ลงนามในใบจ่ายยาและเก็บเอกสารไว้เป็นหลักฐานอ้างอิง")
]

for title_text, detail_text in steps:
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    set_font(run, size=18, bold=True)
    run.font.color.rgb = RGBColor(0, 112, 192) # Blue color for headings
    
    p = doc.add_paragraph()
    run = p.add_run(detail_text)
    set_font(run, size=16)
    doc.add_paragraph()

# ท้ายหน้า
footer = doc.sections[0].footer
p_footer = footer.paragraphs[0]
p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p_footer.add_run("คู่มือฉบับปรับปรุงโดย B-WORD (AI Document Expert)")
set_font(run, size=12)

# Save the document
output_path = '/home/nn/Desktop/งานโรงพยาบาล/wny_hospital/คู่มือขั้นตอนการแพ็คยา_OPD_IPD_ปรับปรุงใหม่.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
